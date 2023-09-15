from docx import Document

company = input('Enter company name: ')
sector = input('Enter sector, you may choose between: tech, finance, insurance, banking, data, consulting: ')
language = input('Enter language, you may choose between nor and eng: ')

if language == 'nor':
    # Load Norwegian template
    document = Document('/Users/jakobalexandersen/job-search/sokndad_layout.docx')
    
    # Replace company name
    for paragraph in document.paragraphs:
        if '[Bedrift]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[Bedrift]', company)

    # Replace sector
    for paragraph in document.paragraphs:
        if '[Industri]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[Industri]', sector)
    
    # Check sector and replace [activity] accordingly
    if sector == 'tech':
        activity_text = 'programvareutvikling og research'
    elif sector in ['finance', 'banking']:
        activity_text = 'finansiell analyse og research'
    elif sector in ['insurance', 'data']:
        activity_text = 'dataanalyse og visualisering'
    elif sector == 'consulting':
        activity_text = 'konsulentvirksomhet og markedsanalyse'
    
    # Replace [Industri] with the sector-specific text
    for paragraph in document.paragraphs:
        if '[activity]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[activity]', activity_text)

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            for elem in run._r:
                if elem.tag.endswith('rFonts'):
                    elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', 'Times New Roman')
                    elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', 'Times New Roman')
                    elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 'Times New Roman')



    

    document.save(f'{company}.docx')


if language == 'eng':
    document = Document('/Users/jakobalexandersen/job-search/english_template.docx')

    for paragraph in document.paragraphs:
        if '[Company]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[Company]', company)
            #set font to times new roman
            paragraph.style.font.name = 'Times New Roman'

    for paragraph in document.paragraphs:
        if '[Sector]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[Sector]', sector)

    if sector == 'tech':
        activity_text = 'software development and research'
    elif sector in ['finance', 'banking']:
        activity_text = 'financial analysis and research'
    elif sector in ['insurance', 'data']:
        activity_text = 'data analysis and visualization'
    elif sector == 'consulting':
        activity_text = 'consulting and market analysis'

    for paragraph in document.paragraphs:
        if '[activity]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[activity]', activity_text)

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            for elem in run._r:
                if elem.tag.endswith('rFonts'):
                    elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', 'Times New Roman')
                    elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', 'Times New Roman')
                    elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 'Times New Roman')

    document.save(f'{company}.docx')

